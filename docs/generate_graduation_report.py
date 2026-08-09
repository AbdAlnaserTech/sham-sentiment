"""
توليد تقرير مشروع التخرج (Word) — تنسيق جامعة الشام.

التشغيل: python docs/generate_graduation_report.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "تقرير_مشروع_التخرج_تحليل_المشاعر.docx"

PROJECT_TITLE_AR = "تحليل مشاعر التعليقات متعددة اللغات"
PROJECT_TITLE_EN = "Multilingual Comment Sentiment Analysis Platform"
DEPARTMENT_AR = "قسم الهندسة المعلوماتية"
YEAR = "2026"
HIJRI_YEAR = "1447"

STUDENTS = [
    "عبد الناصر الحسون",
    "عدي نجار",
    "حسن البكور",
    "محمود السيد",
    "عدنان ناقوح",
]
SUPERVISOR = "د. محمد أسامة"


# ── دوال مساعدة لتنسيق المستند ───────────────────────────────────────────────

def _rtl(p) -> None:
    """محاذاة فقرة من اليمين لليسار."""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p._p.get_or_add_pPr().append(OxmlElement("w:bidi"))


def _center(p) -> None:
    """محاذاة فقرة في المنتصف."""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _font(run, size=14, bold=False, name="Traditional Arabic") -> None:
    """تطبيق خط عربي وحجم وسمك على run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)


def _para(doc, text, *, size=14, bold=False, center=False) -> None:
    """إضافة فقرة نصية."""
    p = doc.add_paragraph()
    ( _center if center else _rtl)(p)
    _font(p.add_run(text), size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5


def _heading(doc, text, level=1) -> None:
    """إضافة عنوان بمستوى محدد."""
    p = doc.add_paragraph()
    _rtl(p)
    _font(p.add_run(text), size={1: 16, 2: 15, 3: 14}.get(level, 14), bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _bullet(doc, text) -> None:
    """إضافة نقطة في قائمة نقطية."""
    p = doc.add_paragraph(style="List Bullet")
    _rtl(p)
    _font(p.add_run(text), size=12)


def _table(doc, headers, rows) -> None:
    """إضافة جدول بعناوين وصفوف."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for par in c.paragraphs:
            _rtl(par)
            for run in par.runs:
                _font(run, size=11, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for par in c.paragraphs:
                _rtl(par)
                for run in par.runs:
                    _font(run, size=11)
    doc.add_paragraph()


def _placeholder_figure(doc, caption: str) -> None:
    """إضافة placeholder لشكل/مخطط يُضاف لاحقاً."""
    _para(doc, f"[ {caption} — أضف الشكل/المخطط هنا ]", center=True, size=11)


def _page_break(doc) -> None:
    """إدراج فاصل صفحة."""
    doc.add_page_break()


def _setup_margins(doc) -> None:
    """ضبط هوامش الصفحة."""
    for s in doc.sections:
        s.top_margin = Mm(40)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(25)
        s.right_margin = Mm(50)


# ── الصفحات التمهيدية ──────────────────────────────────────────────────────────

def _title_page(doc) -> None:
    """صفحة الغلاف."""
    lines = [
        ("الجمهورية العربية السورية", False),
        ("", False),
        ("وزارة التعليم العالي", False),
        ("", False),
        ("جامعة الشام", False),
        ("", False),
        (f"كلية الهندسة — {DEPARTMENT_AR}", False),
        ("", False),
        ("", False),
        (PROJECT_TITLE_AR, True),
        (PROJECT_TITLE_EN, False),
        ("", False),
        ("مشروع تخرج أُعدّ لنيل درجة الإجازة في الهندسة المعلوماتية", False),
        ("", False),
        ("الطلاب", True),
    ]
    for text, bold in lines:
        if text:
            _para(doc, text, center=True, bold=bold)
    for name in STUDENTS:
        _para(doc, name, center=True)
    _para(doc, "", center=True)
    _para(doc, "إشراف", center=True, bold=True)
    _para(doc, SUPERVISOR, center=True)
    _para(doc, f"{YEAR} / {HIJRI_YEAR}", center=True)


def _committee_page(doc) -> None:
    """صفحة تقرير لجنة المناقشة."""
    _page_break(doc)
    _heading(doc, "تقرير لجنة المناقشة")
    _para(
        doc,
        "نؤيد بأننا قرأنا هذا التقرير كلجنة مناقشة وامتحان الطلبة بمحتوياته ونشهد بأنها "
        "كافية كتقرير لمشروع تخرج لنيل درجة الإجازة في تخصص الهندسة المعلوماتية.",
    )
    _para(doc, "")
    _para(doc, "المشرف:                                    رئيس اللجنة:")
    _para(doc, f"الاسم: {SUPERVISOR}              الاسم: ........................")
    _para(doc, "التوقيع: ......................          التوقيع: ......................")
    _para(doc, "التاريخ:    /    /                        التاريخ:    /    /")
    _para(doc, "")
    _para(doc, "الممتحن:")
    _para(doc, "الاسم: ........................")
    _para(doc, "التوقيع: ......................")
    _para(doc, "التاريخ:    /    /")


def _quran_page(doc) -> None:
    """صفحة الآية القرآنية."""
    _page_break(doc)
    _para(doc, "قال تعالى", center=True, bold=True)
    _para(doc, "")
    _para(
        doc,
        "{ يَرْفَعِ اللَّهُ الَّذِينَ آمَنُوا مِنكُمْ وَالَّذِينَ أُوتُوا الْعِلْمَ دَرَجَاتٍ }",
        center=True,
    )
    _para(doc, "سورة المجادلة — آية 11", center=True)


def _thanks_page(doc) -> None:
    """صفحة الشكر والتقدير."""
    _page_break(doc)
    _heading(doc, "شكر وتقدير")
    _para(
        doc,
        f"بعد رحلة من البحث والعمل الجماعي، نتوجه بأسمى عبارات الشكر والتقدير إلى "
        f"مشرف مشروعنا {SUPERVISOR} على إشرافه وتوجيهاته القيّمة، ولما قدمه من دعم "
        "علمي وعملي طوال مراحل التطوير والاختبار.",
    )
    _para(
        doc,
        "نشكر أساتذة قسم الهندسة المعلوماتية في جامعة الشام، ومركز الحاسوب، "
        "وكل من ساهم بملاحظاته أو دعمه التقني.",
    )
    _para(doc, "فريق العمل:", bold=True)
    for n in STUDENTS:
        _bullet(doc, n)


def _dedication_page(doc) -> None:
    """صفحة الإهداء."""
    _page_break(doc)
    _heading(doc, "الإهداء")
    _para(doc, "إلى من علّمنا أن العلم نور...", center=True)
    _para(doc, "")
    _para(doc, "(( صفحة الإهداء — يُعدّلها الفريق: والدين، أم، أهل، ... ))", center=True, size=11)


def _abstract_page(doc) -> None:
    """صفحة الخلاصة (عربي + Abstract إنجليزي)."""
    _page_break(doc)
    _heading(doc, "الخلاصة")
    for t in (
        "تزداد أهمية تحليل مشاعر التعليقات مع نمو التجارة الإلكترونية ومنصات التواصل. "
        "يواجه التحليل تحديات عند التعامل مع العربية الفصحى واللهجة الشامية والإنجليزية "
        "في آن واحد. يقدّم هذا المشروع منصة ويب لتحليل التعليقات وتصنيفها: إيجابي، محايد، سلبي.",
        "يعتمد النظام على BERT (XLM-RoBERTa) كنموذج افتراضي، مع مسار TF-IDF للمقارنة "
        "والتدريب. تشمل المنصة واجهة Streamlit، API عبر FastAPI، قاعدة SQLite، جلب "
        "تعليقات من YouTube وGoogle Play، وتصدير PDF/Excel.",
        "على 513 تعليقاً في validation_comments.csv: BERT حقق دقة 65.3% وMacro-F1=0.63، "
        "متفوقاً على TF-IDF (53.8% / 0.53). أفضل أداء في الإيجابي والسلبي؛ المحايد "
        "يحتاج بيانات إضافية. المشروع جاهز للعرض مع 43 اختباراً ناجحاً.",
    ):
        _para(doc, t)
    _page_break(doc)
    _heading(doc, "Abstract")
    for t in (
        "This graduation project implements a multilingual sentiment analysis platform for "
        "English, Modern Standard Arabic, and Levantine Arabic comments.",
        "The stack includes BERT/XLM-RoBERTa, TF-IDF baseline, Streamlit UI, FastAPI, "
        "and SQLite analytics storage.",
        "Validation: BERT 65.3% accuracy vs TF-IDF 53.8%. Future: expand Shami data, "
        "cloud deployment.",
    ):
        _para(doc, t)


def _symbols_page(doc) -> None:
    """قائمة الرموز والاختصارات."""
    _page_break(doc)
    _heading(doc, "قائمة بالرموز")
    _table(
        doc,
        ["الرمز", "المعنى"],
        [
            ("NLP", "Natural Language Processing — معالجة اللغة الطبيعية"),
            ("BERT", "Bidirectional Encoder Representations from Transformers"),
            ("TF-IDF", "Term Frequency – Inverse Document Frequency"),
            ("F1", "مقياس F1-Score"),
            ("API", "Application Programming Interface"),
            ("RTL", "Right-to-Left — اتجاه من اليمين لليسار"),
        ],
    )


def _toc_and_lists(doc) -> None:
    """الفهرس وقوائم الأشكال والجداول."""
    _page_break(doc)
    _heading(doc, "الفهرس")
    entries = [
        "سورة المجادلة ................. i",
        "شكر وتقدير .................... ii",
        "الإهداء ....................... iii",
        "الخلاصة ....................... iv",
        "Abstract ...................... v",
        "الفهرس ........................ vi",
        "قائمة بالأشكال ................ vii",
        "قائمة بالجداول ................ viii",
        "قائمة بالرموز ................. ix",
        "",
        "الفصل الأول: المقدمة",
        "  1.1 المقدمة",
        "  1.2 نظرة عامة على النظام",
        "  1.3 الدراسات السابقة",
        "  1.4 الفكرة العامة",
        "  1.5 المشكلة",
        "  1.6 الأهداف",
        "  1.7 الأدوار",
        "  1.8 المنهجية (SDLC)",
        "  1.9 الأهمية",
        "  1.10 نطاق المشروع",
        "  1.11 نقاط القوة",
        "  1.12 الجدول الزمني",
        "",
        "الفصل الثاني: تحليل النظام",
        "  2.1–2.8 المتطلبات والجدوى وUse Cases",
        "الفصل الثالث: التصميم",
        "  3.1–3.6 قاعدة البيانات والمعمارية",
        "الفصل الرابع: التنفيذ",
        "  4.1–4.7 النماذج والواجهة والAPI",
        "الفصل الخامس: النتائج",
        "  5.1–5.5 التقييم والتوصيات",
        "المصادر · الملاحق",
    ]
    for e in entries:
        _para(doc, e, size=12)

    _page_break(doc)
    _heading(doc, "قائمة بالأشكال")
    for cap in (
        "(1-1) مخطط هرمي للنظام",
        "(1-2) مخطط تدفق تحليل التعليق",
        "(2-1) مخطط حالات الاستخدام UML",
        "(2-2) مخطط النشاط — تحليل دفعة",
        "(3-1) ERD قاعدة البيانات",
        "(3-2) معمارية الطبقات",
        "(4-1) شاشة تسجيل الدخول",
        "(4-2) تحليل تعليق واحد",
        "(4-3) لوحة التحكم",
        "(5-1) Confusion Matrix",
    ):
        _para(doc, f"الشكل {cap}", size=12)

    _page_break(doc)
    _heading(doc, "قائمة بالجداول")
    for cap in (
        "(1) مقارنة TF-IDF و BERT",
        "(2) نتائج BERT التفصيلية",
        "(3) أدوار المستخدمين",
        "(4) المتطلبات الوظيفية",
        "(5) الجدول الزمني Gantt",
        "(6) هيكل مجلدات المشروع",
    ):
        _para(doc, f"الجدول {cap}", size=12)


# ── الفصول ──────────────────────────────────────────────────────────────

def _chapter1(doc) -> None:
    """الفصل الأول: المقدمة."""
    _page_break(doc)
    _heading(doc, "الفصل الأول: المقدمة")

    _heading(doc, "1.1 المقدمة", 2)
    _para(
        doc,
        "في عصر البيانات الضخمة، تمثل التعليقات النصية مصدراً غنياً لآراء العملاء "
        "والجمهور. يندرج تحليل المشاعر (Sentiment Analysis) ضمن معالجة اللغة الطبيعية "
        "(NLP) ويهدف إلى استخراج الانطباع العاطفي من النص. يقدّم هذا التقرير مشروع "
        "تخرج لبناء منصة عملية تخدم جامعة الشام — قسم الهندسة المعلوماتية.",
    )

    _heading(doc, "1.2 نظرة عامة على النظام", 2)
    _para(
        doc,
        "المنصة طبقة عرض (Streamlit) + نواة Python (src/) + نماذج ML + SQLite. "
        "المستخدم يسجّل الدخول، يختار تحليل فردي/مجموعة/جلب حي، ويرى النتائج مع "
        "ثقة التنبؤ وكلمات مفتاحية وتنبيهات.",
    )
    _placeholder_figure(doc, "الشكل 1-1: مخطط هرمي عام")

    _heading(doc, "1.3 الدراسات السابقة", 2)
    _para(
        doc,
        "• Pang & Lee (2008): أسس تحليل المشاعر الإنجليزي.\n"
        "• ASTD: مجموعة tweets عربية للمشاعر (Nabil et al.).\n"
        "• Devlin et al.: BERT — تحولات Transformer.\n"
        "• Conneau et al.: XLM-R — متعدد اللغات.\n"
        "• CAMeLBERT: نماذج عربية متخصصة.\n"
        "يُميّز مشروعنا بدمج واجهة عربية + دعm شامي + منصة متكاملة للعرض.",
    )

    _heading(doc, "1.4 الفكرة العامة للمشروع", 2)
    _para(
        doc,
        "منصة «تحليل آراء العملاء» — إدخال تعليق أو ملف CSV أو رابط YouTube/Play، "
        "كشف اللغة تلقائياً، تصنيف المشاعر، حفظ في DB، عرض analytics وتصدير.",
    )

    _heading(doc, "1.5 المشكلة", 2)
    _bullet(doc, "غياب أدوات محلية تدعم العربية والشامية معاً.")
    _bullet(doc, "صعوبة تحليل آلاف التعليقات يدوياً.")
    _bullet(doc, "حاجة المؤسسات التعليمية لمشاريع NLP تطبيقية.")

    _heading(doc, "1.6 أهداف المشروع", 2)
    goals = [
        "تصنيف ثلاثي: positive / neutral / negative.",
        "دعم en, ar_fusha, ar_shami.",
        "واجهة Streamlit + صلاحيات مستخدمين.",
        "تقييم BERT vs TF-IDF على بيانات تحقق.",
        "توثيق وتجهيز للعرض والنشر.",
    ]
    for i, g in enumerate(goals, 1):
        _bullet(doc, f"1.6.{i} {g}")

    _heading(doc, "1.7 الأدوار الرئيسية في النظام", 2)
    _table(
        doc,
        ["الدور", "الوصف"],
        [
            ("admin", "مدير — حذف دفعات، إدارة كاملة"),
            ("analyst", "محلل — تحليل + لوحة"),
            ("viewer", "عرض — لوحة وحول فقط"),
        ],
    )
    _para(doc, "الجدول (3): أدوار المستخدمين.")

    _heading(doc, "1.8 المنهجية (SDLC)", 2)
    _para(
        doc,
        "اتُبعت دورة حياة النظام: (1) التخطيط، (2) جمع المتطلبات، (3) التحليل، "
        "(4) التصميم، (5) التنفيذ، (6) الاختبار pytest، (7) النشر التجريبي Streamlit، "
        "(8) التوثيق. أُنجز التطوير بشكل تكراري (Agile خفيف) مع مراجعات أسبوعية للمشرف.",
    )

    _heading(doc, "1.9 أهمية المشروع", 2)
    _bullet(doc, "للطالب: تطبيق AI/NLP في مشروع تخرج.")
    _bullet(doc, "للشركات: فهم رضا العملاء.")
    _bullet(doc, "للجامعة: نموذج قابل للنشر على Streamlit Cloud.")

    _heading(doc, "1.10 نطاق المشروع", 2)
    _para(doc, "داخل النطاق: NLP، واجهة، API، DB، YouTube/Play.")
    _para(doc, "خارج النطاق: OAuth، Kubernetes، Reddit API (غير مستقر).")

    _heading(doc, "1.11 من أهم نقاط القوة", 2)
    _bullet(doc, "دعم ثلاث لغات/أنماط.")
    _bullet(doc, "BERT + baseline للمقارنة الأكاديمية.")
    _bullet(doc, "43 unit test.")
    _bullet(doc, "واجهة عربية RTL.")

    _heading(doc, "1.12 الجدول الزمني", 2)
    _table(
        doc,
        ["المهمة", "الأسابيع"],
        [
            ("التخطيط واختيار الفكرة", "3"),
            ("جمع المتطلبات ودراسة الأعمال السابقة", "4"),
            ("تحليل وتصميم", "6"),
            ("تطوير النماذج", "6"),
            ("تطوير الواجهة وDB", "5"),
            ("اختبار وتوثيق", "4"),
            ("الإجمالي التقريبي", "28"),
        ],
    )
    _para(doc, "الجدول (5): الجدول الزمني. [ أضف مخطط Gantt في Word ]")
    _placeholder_figure(doc, "الشكل 1-2: Gantt Chart")


def _chapter2(doc) -> None:
    """الفصل الثاني: تحليل النظام."""
    _page_break(doc)
    _heading(doc, "الفصل الثاني: تحليل النظام")

    _heading(doc, "2.1 المقدمة", 2)
    _para(doc, "تحليل المشكلة، المتطلبات، الجدوى، وحالات الاستخدام.")

    _heading(doc, "2.2 استكشاف المشكلة", 2)
    _para(
        doc,
        "العملاء يكتبون تعليقات متنوعة اللغة والأسلوب. التحليل اليدوي بطيء. "
        "الحلول السحابية الأجنبية قد لا تفهم اللهجة الشامية. نحتاج نظاماً محلياً "
        "قابلاً للتخصيص.",
    )

    _heading(doc, "2.3 تحديد احتياجات المستخدم", 2)
    _bullet(doc, "إدخال سريع لتعليق واحد.")
    _bullet(doc, "رفع CSV لآلاف التعليقات.")
    _bullet(doc, "تقرير PDF للجنة.")
    _bullet(doc, "تنبيه عند ارتفاع السلبية.")

    _heading(doc, "2.4 المتطلبات الوظيفية", 2)
    _table(
        doc,
        ["#", "المتطلب", "الأولوية"],
        [
            ("FR1", "تسجيل دخول وصلاحيات", "عالية"),
            ("FR2", "تحليل تعليق واحد", "عالية"),
            ("FR3", "تحليل دفعة CSV", "عالية"),
            ("FR4", "جلب YouTube/Play", "متوسطة"),
            ("FR5", "لوحة تحكم KPIs", "عالية"),
            ("FR6", "تصدير PDF/Excel", "متوسطة"),
            ("FR7", "حفظ الدفعات SQLite", "متوسطة"),
        ],
    )
    _para(doc, "الجدول (4): المتطلبات الوظيفية.")

    _heading(doc, "2.5 المتطلبات غير الوظيفية", 2)
    _bullet(doc, "NFR1: أمان كلمات المرور PBKDF2.")
    _bullet(doc, "NFR2: زمن تحميل BERT ≤30s أول مرة (cache).")
    _bullet(doc, "NFR3: واجهة RTL.")
    _bullet(doc, "NFR4: max_text_length = 5000.")

    _heading(doc, "2.6 دراسة الجدوى التقنية", 2)
    _para(doc, "Python ecosystem ناضج. Hugging Face يوفّر XLM-R. Streamlit يسرّع UI. الجدوى: عالية.")

    _heading(doc, "2.7 دراسة الجدوى الاقتصادية", 2)
    _para(doc, "مفتوح المصدر. تكلفة Streamlit Cloud مجانية للعرض. لا تراخيص.")

    _heading(doc, "2.8 نمذجة النظام — Use Cases", 2)
    _para(doc, "Actors: Analyst, Admin, Viewer, External API.")
    _bullet(doc, "UC1: Login")
    _bullet(doc, "UC2: Analyze Single Comment")
    _bullet(doc, "UC3: Batch Upload")
    _bullet(doc, "UC4: Live Import")
    _bullet(doc, "UC5: View Dashboard")
    _placeholder_figure(doc, "الشكل 2-1: Use Case Diagram")


def _chapter3(doc) -> None:
    """الفصل الثالث: التحليل والتصميم."""
    _page_break(doc)
    _heading(doc, "الفصل الثالث: مرحلة التحليل والتصميم")

    _heading(doc, "3.1 المقدمة", 2)
    _para(doc, "تصميم البيانات، المعمارية، وتدفق المعالجة.")

    _heading(doc, "3.2 المدخلات", 2)
    _bullet(doc, "نص UTF-8 حتى 5000 حرف.")
    _bullet(doc, "CSV: عمود text (+ language اختياري).")
    _bullet(doc, "URL YouTube / Google Play.")

    _heading(doc, "3.3 المخرجات", 2)
    _bullet(doc, "sentiment, confidence, language, is_reliable.")
    _bullet(doc, "keywords حسب المشاعر.")
    _bullet(doc, "alerts (negative_spike, low_confidence).")

    _heading(doc, "3.4 مخطط تدفق التحليل", 2)
    _para(
        doc,
        "تعليق → كشف لغة → preprocessing → BERT/TF-IDF → نتيجة + ثقة → "
        "analytics → (اختياري) حفظ DB.",
    )
    _placeholder_figure(doc, "الشكل 1-2: Activity Diagram")

    _heading(doc, "3.5 تصميم قاعدة البيانات", 2)
    _table(
        doc,
        ["الجدول", "الحقول الرئيسية"],
        [
            ("users", "id, username, password_hash, role"),
            ("analysis_batches", "id, user_id, source, total, avg_confidence"),
            ("analysis_items", "batch_id, text, sentiment, confidence, language"),
            ("alerts", "batch_id, alert_type, severity, message"),
        ],
    )
    _placeholder_figure(doc, "الشكل 3-1: ERD")

    _heading(doc, "3.6 معمارية النظام", 2)
    _para(
        doc,
        "Presentation: app/main.py, components/*\n"
        "Business: src/models, src/analytics, src/preprocessing\n"
        "Data: src/db, data/*.csv\n"
        "Integration: src/api/server.py, comment_fetcher.py",
    )
    _placeholder_figure(doc, "الشكل 3-2: Layered Architecture")


def _chapter4(doc) -> None:
    """الفصل الرابع: التنفيذ."""
    _page_break(doc)
    _heading(doc, "الفصل الرابع: مرحلة التنفيذ")

    _heading(doc, "4.1 المقدمة", 2)
    _para(doc, "تفاصيل التطبيق البرمجي.")

    _heading(doc, "4.2 هيكل المشروع", 2)
    _table(
        doc,
        ["المجلد", "الوظيفة"],
        [
            ("app/", "واجهة Streamlit"),
            ("src/models/", "BERT, TF-IDF, registry"),
            ("src/language.py", "كشف فصحى/شامي"),
            ("src/preprocessing.py", "تطبيع عربي"),
            ("src/db/", "SQLite repository"),
            ("tests/", "43 pytest"),
        ],
    )
    _para(doc, "الجدول (6): هيكل المجلدات.")

    _heading(doc, "4.3 التقنيات", 2)
    _table(
        doc,
        ["التقنية", "الاستخدام"],
        [
            ("Python 3.11", "لغة أساسية"),
            ("Streamlit", "UI"),
            ("FastAPI", "REST API"),
            ("transformers", "BERT"),
            ("scikit-learn", "TF-IDF"),
            ("SQLite", "تخزين"),
        ],
    )

    _heading(doc, "4.4 تنفيذ كشف اللغة", 2)
    _para(
        doc,
        "language.py: detect_language() — قواعد SHAMI_HINT_WORDS للتمييز بين "
        "ar_shami و ar_fusha، regex للاتجاه الإنجليزي.",
    )

    _heading(doc, "4.5 تنفيذ النماذج", 2)
    _para(doc, "TF-IDF: trainer.py — GridSearchCV + CalibratedClassifierCV.")
    _para(doc, "BERT: bert_predictor.py — XLM-RoBERTa ensemble، fallback models.")
    _para(doc, "registry.py — load_predictor('bert'|'tfidf').")

    _heading(doc, "4.6 تنفيذ الواجهة", 2)
    _para(
        doc,
        "5 تبويبات: لوحة التحكم، تعليق واحد، مجموعة، جلب حي، حول المشروع. "
        "shared.py — cache BERT. auth_panel — PBKDF2 login.",
    )
    _placeholder_figure(doc, "الشكل 4-1: Login")
    _placeholder_figure(doc, "الشكل 4-2: Single Analysis")
    _placeholder_figure(doc, "الشكل 4-3: Dashboard")

    _heading(doc, "4.7 API", 2)
    _para(doc, "POST /api/v1/analyze, /batch — optional X-API-Key header.")


def _chapter5(doc) -> None:
    """الفصل الخامس: النتائج والتوصيات."""
    _page_break(doc)
    _heading(doc, "الفصل الخامس: النتائج والتوصيات")

    _heading(doc, "5.1 منهجية التقييم", 2)
    _para(
        doc,
        "Dataset: data/real/validation_comments.csv — 513 تعليقاً. "
        "Metrics: Accuracy, Precision, Recall, F1-macro. مقارنة bert vs tfidf.",
    )

    _heading(doc, "5.2 نتائج المقارنة", 2)
    _table(
        doc,
        ["النموذج", "Accuracy", "Macro-F1"],
        [
            ("TF-IDF + LogReg", "53.8%", "0.53"),
            ("BERT (XLM-R)", "65.3%", "0.63"),
        ],
    )
    _para(doc, "الجدول (1): مقارنة النماذج.")

    _heading(doc, "5.3 نتائج BERT التفصيلية", 2)
    _table(
        doc,
        ["الفئة", "Precision", "Recall", "F1"],
        [
            ("negative", "0.65", "0.80", "0.72"),
            ("neutral", "0.61", "0.35", "0.44"),
            ("positive", "0.67", "0.81", "0.74"),
        ],
    )
    _para(doc, "الجدول (2): classification report — BERT.")

    _heading(doc, "5.4 مناقشة", 2)
    _para(
        doc,
        "BERT أفضل بـ ~11 نقطة accuracy. المحايد الأضعف — تداخل lexical. "
        "عينة shami=16 فقط — لا تعميم قوي. 43 test passed.",
    )
    _placeholder_figure(doc, "الشكل 5-1: Confusion Matrix — [أضف من evaluate.py]")

    _heading(doc, "5.5 الاستنتاجات والعمل المستقبلي", 2)
    _bullet(doc, "تم تحقيق منصة متكاملة.")
    _bullet(doc, "BERT مناسب للعرض.")
    _bullet(doc, "مستقبلاً: بيانات شami أكثر، fine-tune، PostgreSQL، Docker.")

    _heading(doc, "5.6 الاختبار", 2)
    _para(doc, "pytest tests/ — 43 passed. يغطي preprocessing, language, API, export.")


def _references(doc) -> None:
    """قسم المصادر والمراجع."""
    _page_break(doc)
    _heading(doc, "المصادر")
    refs = [
        "Pang, B., & Lee, L. (2008). Opinion Mining and Sentiment Analysis.",
        "Devlin, J., et al. (2019). BERT: Pre-training of Deep Bidirectional Transformers.",
        "Conneau, A., et al. (2020). Unsupervised Cross-lingual Representation Learning at Scale (XLM-R).",
        "Nabil, M., et al. ASTD: Arabic Sentiment Tweets Dataset.",
        "Inoue, G., et al. CAMeLBERT: Transformer-based Arabic Models.",
        "Pedregosa, F., et al. Scikit-learn: Machine Learning in Python.",
        "Streamlit Inc. Streamlit Documentation — https://streamlit.io",
        "FastAPI — https://fastapi.tiangolo.com",
        "Hugging Face Transformers — https://huggingface.co/docs/transformers",
    ]
    for i, r in enumerate(refs, 1):
        _para(doc, f"[{i}] {r}", size=12)


def _appendices(doc) -> None:
    """الملاحق — لقطات الشاشة وأوامر التشغيل."""
    _page_break(doc)
    _heading(doc, "الملحق أ — لقطات الشاشة")
    _para(doc, "أضف 4–6 screenshots:")
    for s in (
        "تسجيل الدخول",
        "تحليل تعليق (إيجابي/سلبي)",
        "نتائج دفعة + رسم pie",
        "لوحة التحكم",
        "تبويب حول المشروع",
        "جلب تعليقات YouTube",
    ):
        _bullet(doc, s)
        _placeholder_figure(doc, f"Screenshot: {s}")

    _page_break(doc)
    _heading(doc, "الملحق ب — أوامر التشغيل")
    _para(doc, "pip install -r requirements.txt")
    _para(doc, "python -m streamlit run app/main.py")
    _para(doc, "python launchers/evaluate.py --data data/real/validation_comments.csv --model bert")


def build_report() -> Path:
    """بناء التقرير الكامل وحفظه كملف Word."""
    doc = Document()
    _setup_margins(doc)
    _title_page(doc)
    _committee_page(doc)
    _quran_page(doc)
    _thanks_page(doc)
    _dedication_page(doc)
    _abstract_page(doc)
    _toc_and_lists(doc)
    _symbols_page(doc)
    _chapter1(doc)
    _chapter2(doc)
    _chapter3(doc)
    _chapter4(doc)
    _chapter5(doc)
    _references(doc)
    _appendices(doc)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    return OUT_PATH


if __name__ == "__main__":
    print(f"Report saved: {build_report()}")
